#!/usr/bin/env python3
"""
Universal File Extractor for Elasticsearch Batch Processing
===========================================================
Supports: PDF, TXT, DOC, DOCX, XLS, XLSX, CSV, HTML, XML, JSON, MD
"""

import sys
import json
import logging
from pathlib import Path
from typing import Tuple, Dict, Optional, List
import hashlib
from datetime import datetime

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('UniversalExtractor')

class UniversalFileExtractor:
    """Extract text and metadata from various file formats"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self.extract_pdf,
            '.txt': self.extract_text_file,
            '.md': self.extract_text_file,
            '.doc': self.extract_doc,
            '.docx': self.extract_docx,
            '.xls': self.extract_xls,
            '.xlsx': self.extract_xlsx,
            '.csv': self.extract_csv,
            '.html': self.extract_html,
            '.htm': self.extract_html,
            '.xml': self.extract_xml,
            '.json': self.extract_json,
            '.log': self.extract_text_file,
            '.rtf': self.extract_rtf,
            '.odt': self.extract_odt,
            '.ods': self.extract_ods,
        }
        
        self.stats = {
            'processed': 0,
            'successful': 0,
            'failed': 0,
            'by_type': {}
        }
        
        # Check available libraries
        self._check_dependencies()
        
    def _check_dependencies(self):
        """Check which extraction libraries are available"""
        self.available_libs = {}
        
        try:
            import fitz
            self.available_libs['pymupdf'] = True
        except ImportError:
            self.available_libs['pymupdf'] = False
            logger.warning("PyMuPDF not installed - PDF support limited")
            
        try:
            import docx
            self.available_libs['python-docx'] = True
        except ImportError:
            self.available_libs['python-docx'] = False
            logger.warning("python-docx not installed - DOCX support disabled")
            
        try:
            import openpyxl
            self.available_libs['openpyxl'] = True
        except ImportError:
            self.available_libs['openpyxl'] = False
            logger.warning("openpyxl not installed - XLSX support disabled")
            
        try:
            import pandas
            self.available_libs['pandas'] = True
        except ImportError:
            self.available_libs['pandas'] = False
            logger.warning("pandas not installed - Advanced spreadsheet support disabled")
            
        try:
            from bs4 import BeautifulSoup
            self.available_libs['beautifulsoup4'] = True
        except ImportError:
            self.available_libs['beautifulsoup4'] = False
            logger.warning("BeautifulSoup not installed - HTML parsing limited")
            
    def extract(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text and metadata from any supported file"""
        
        if not file_path.exists():
            return "", {"error": "File not found"}
            
        # Get file extension
        ext = file_path.suffix.lower()
        
        # Update stats
        self.stats['processed'] += 1
        self.stats['by_type'][ext] = self.stats['by_type'].get(ext, 0) + 1
        
        # Get appropriate extractor
        extractor = self.supported_formats.get(ext)
        
        if not extractor:
            logger.warning(f"Unsupported file type: {ext}")
            self.stats['failed'] += 1
            return f"[Unsupported file type: {ext}]", {"error": f"No extractor for {ext}"}
            
        try:
            # Extract content
            text, metadata = extractor(file_path)
            
            # Add common metadata
            metadata.update({
                'file_type': ext,
                'file_size': file_path.stat().st_size,
                'file_name': file_path.name,
                'extraction_date': datetime.now().isoformat(),
                'file_hash': self._calculate_file_hash(file_path)
            })
            
            self.stats['successful'] += 1
            return text, metadata
            
        except Exception as e:
            logger.error(f"Extraction failed for {file_path.name}: {e}")
            self.stats['failed'] += 1
            return f"[Extraction failed: {str(e)}]", {"error": str(e)}
            
    def extract_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from PDF files"""
        if not self.available_libs.get('pymupdf'):
            return self._fallback_pdf_extraction(file_path)
            
        import fitz
        
        doc = fitz.open(str(file_path))
        text_parts = []
        metadata = {
            'page_count': len(doc),
            'title': doc.metadata.get('title', ''),
            'author': doc.metadata.get('author', ''),
            'subject': doc.metadata.get('subject', ''),
            'creation_date': str(doc.metadata.get('creationDate', '')),
        }
        
        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                
        doc.close()
        
        full_text = '\n\n'.join(text_parts) if text_parts else "[No extractable text - OCR may be needed]"
        metadata['has_text'] = bool(text_parts)
        
        return full_text, metadata
        
    def extract_text_file(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract content from plain text files (TXT, MD, LOG)"""
        encodings = ['utf-8', 'iso-8859-1', 'windows-1252', 'ascii']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                    
                metadata = {
                    'encoding': encoding,
                    'line_count': text.count('\n') + 1,
                    'word_count': len(text.split()),
                    'char_count': len(text)
                }
                
                return text, metadata
            except UnicodeDecodeError:
                continue
                
        return "[Unable to decode file - unknown encoding]", {"error": "Encoding detection failed"}
        
    def extract_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from DOCX files"""
        if not self.available_libs.get('python-docx'):
            return "[python-docx required for DOCX files]", {"error": "Missing dependency"}
            
        import docx
        
        doc = docx.Document(str(file_path))
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
                
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    table_data.append(' | '.join(row_data))
            if table_data:
                tables_text.append('\n'.join(table_data))
                
        # Combine all text
        all_text = '\n\n'.join(paragraphs)
        if tables_text:
            all_text += '\n\n--- Tables ---\n\n' + '\n\n'.join(tables_text)
            
        metadata = {
            'paragraph_count': len(paragraphs),
            'table_count': len(doc.tables),
            'has_images': bool(doc.inline_shapes),
            'core_properties': {
                'author': doc.core_properties.author or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                'title': doc.core_properties.title or '',
                'subject': doc.core_properties.subject or '',
            }
        }
        
        return all_text, metadata
        
    def extract_doc(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from legacy DOC files"""
        # Try using python-docx2txt or fallback
        try:
            import docx2txt
            text = docx2txt.process(str(file_path))
            return text, {"extraction_method": "docx2txt"}
        except ImportError:
            # Fallback: try to convert using system tools
            return self._fallback_doc_extraction(file_path)
            
    def extract_xlsx(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract data from XLSX files"""
        if self.available_libs.get('pandas'):
            return self._extract_excel_pandas(file_path)
        elif self.available_libs.get('openpyxl'):
            return self._extract_excel_openpyxl(file_path)
        else:
            return "[pandas or openpyxl required for Excel files]", {"error": "Missing dependency"}
            
    def _extract_excel_pandas(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract Excel data using pandas"""
        import pandas as pd
        
        text_parts = []
        metadata = {'sheets': {}}
        
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        
        for sheet_name in excel_file.sheet_names:
            df = excel_file.parse(sheet_name)
            
            # Convert to text
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            text_parts.append(df.to_string())
            
            # Sheet metadata
            metadata['sheets'][sheet_name] = {
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns)
            }
            
        metadata['sheet_count'] = len(excel_file.sheet_names)
        
        return '\n\n'.join(text_parts), metadata
        
    def _extract_excel_openpyxl(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract Excel data using openpyxl"""
        import openpyxl
        
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        text_parts = []
        metadata = {'sheets': {}}
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            
            # Extract cell values
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else '' for cell in row]
                if any(row_data):
                    rows.append('\t'.join(row_data))
                    
            text_parts.append('\n'.join(rows))
            
            metadata['sheets'][sheet_name] = {
                'max_row': ws.max_row,
                'max_column': ws.max_column
            }
            
        wb.close()
        metadata['sheet_count'] = len(wb.sheetnames)
        
        return '\n\n'.join(text_parts), metadata
        
    def extract_xls(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract data from legacy XLS files"""
        if self.available_libs.get('pandas'):
            try:
                return self._extract_excel_pandas(file_path)
            except Exception as e:
                logger.warning(f"Pandas failed for XLS: {e}")
                
        # Fallback for XLS
        return "[XLS extraction requires pandas with xlrd]", {"error": "XLS support limited"}
        
    def extract_csv(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract data from CSV files"""
        import csv
        
        text_parts = []
        metadata = {}
        
        # Detect delimiter
        with open(file_path, 'r', encoding='utf-8-sig') as f:
            sample = f.read(1024)
            sniffer = csv.Sniffer()
            try:
                delimiter = sniffer.sniff(sample).delimiter
            except:
                delimiter = ','
                
        # Read CSV
        with open(file_path, 'r', encoding='utf-8-sig') as f:
            reader = csv.reader(f, delimiter=delimiter)
            rows = list(reader)
            
        if rows:
            # Assume first row is header
            headers = rows[0]
            metadata['columns'] = headers
            metadata['row_count'] = len(rows) - 1
            metadata['delimiter'] = delimiter
            
            # Convert to text
            for row in rows:
                text_parts.append('\t'.join(str(cell) for cell in row))
                
        return '\n'.join(text_parts), metadata
        
    def extract_html(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from HTML files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
            
        if self.available_libs.get('beautifulsoup4'):
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
                
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            metadata = {
                'title': soup.title.string if soup.title else '',
                'links_count': len(soup.find_all('a')),
                'images_count': len(soup.find_all('img')),
                'has_tables': bool(soup.find_all('table'))
            }
            
            return text, metadata
        else:
            # Basic extraction
            import re
            text = re.sub('<[^<]+?>', '', html_content)
            return text, {"extraction_method": "regex"}
            
    def extract_xml(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from XML files"""
        import xml.etree.ElementTree as ET
        
        tree = ET.parse(file_path)
        root = tree.getroot()
        
        # Extract all text content
        text_parts = []
        
        def extract_text_from_element(element, level=0):
            if element.text and element.text.strip():
                indent = "  " * level
                text_parts.append(f"{indent}{element.tag}: {element.text.strip()}")
            for child in element:
                extract_text_from_element(child, level + 1)
                
        extract_text_from_element(root)
        
        metadata = {
            'root_tag': root.tag,
            'element_count': len(root.findall('.//')),
            'namespaces': dict(root.attrib) if root.attrib else {}
        }
        
        return '\n'.join(text_parts), metadata
        
    def extract_json(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract data from JSON files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
        # Pretty print JSON as text
        text = json.dumps(data, indent=2, ensure_ascii=False)
        
        metadata = {
            'type': type(data).__name__,
            'size': len(data) if isinstance(data, (list, dict)) else 1
        }
        
        if isinstance(data, dict):
            metadata['keys'] = list(data.keys())
        elif isinstance(data, list) and data:
            metadata['first_item_type'] = type(data[0]).__name__
            
        return text, metadata
        
    def extract_rtf(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from RTF files"""
        try:
            import striprtf
            with open(file_path, 'r', encoding='utf-8') as f:
                rtf_content = f.read()
            text = striprtf.rtf_to_text(rtf_content)
            return text, {"extraction_method": "striprtf"}
        except ImportError:
            # Basic RTF extraction
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            # Remove RTF control words (very basic)
            import re
            text = re.sub(r'\\[a-z]+\d*\s?', '', content)
            text = re.sub(r'[{}]', '', text)
            return text, {"extraction_method": "regex_fallback"}
            
    def extract_odt(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from OpenDocument Text files"""
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            text_parts = []
            
            with zipfile.ZipFile(file_path, 'r') as z:
                # Read content.xml
                with z.open('content.xml') as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    
                    # Extract text from all text elements
                    for elem in root.iter():
                        if elem.text and elem.text.strip():
                            text_parts.append(elem.text.strip())
                            
            return ' '.join(text_parts), {"extraction_method": "xml_parsing"}
            
        except Exception as e:
            return f"[ODT extraction failed: {e}]", {"error": str(e)}
            
    def extract_ods(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract data from OpenDocument Spreadsheet files"""
        if self.available_libs.get('pandas'):
            try:
                import pandas as pd
                df = pd.read_excel(file_path, engine='odf')
                return df.to_string(), {"extraction_method": "pandas_odf"}
            except:
                pass
                
        # Fallback
        return "[ODS extraction requires pandas with odfpy]", {"error": "Missing dependency"}
        
    def _fallback_pdf_extraction(self, file_path: Path) -> Tuple[str, Dict]:
        """Fallback PDF extraction using system tools"""
        import subprocess
        
        try:
            # Try pdftotext (if available)
            result = subprocess.run(
                ['pdftotext', str(file_path), '-'],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout, {"extraction_method": "pdftotext"}
        except:
            pass
            
        return "[PDF extraction requires PyMuPDF or pdftotext]", {"error": "No PDF extractor available"}
        
    def _fallback_doc_extraction(self, file_path: Path) -> Tuple[str, Dict]:
        """Fallback DOC extraction using system tools"""
        import subprocess
        
        try:
            # Try antiword (if available)
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout, {"extraction_method": "antiword"}
        except:
            pass
            
        return "[DOC extraction requires python-docx or antiword]", {"error": "No DOC extractor available"}
        
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
        
    def get_stats(self) -> Dict:
        """Get extraction statistics"""
        return dict(self.stats)


def install_dependencies():
    """Install recommended dependencies"""
    print("Installing recommended dependencies...")
    
    dependencies = [
        'pymupdf',           # PDF
        'python-docx',       # DOCX
        'openpyxl',         # XLSX
        'pandas',           # Advanced spreadsheets
        'beautifulsoup4',   # HTML
        'lxml',             # XML parsing
        'striprtf',         # RTF
        'odfpy',            # ODF support
        'docx2txt',         # DOC fallback
    ]
    
    import subprocess
    
    for dep in dependencies:
        print(f"Installing {dep}...")
        try:
            subprocess.run([sys.executable, '-m', 'pip', 'install', dep], check=True)
            print(f"✅ {dep} installed")
        except:
            print(f"⚠️  Failed to install {dep}")
            
    print("\nInstallation complete!")


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser()
    parser.add_argument('file', nargs='?', help='File to extract')
    parser.add_argument('--install', action='store_true', help='Install dependencies')
    parser.add_argument('--test', action='store_true', help='Test extraction')
    
    args = parser.parse_args()
    
    if args.install:
        install_dependencies()
        sys.exit(0)
        
    extractor = UniversalFileExtractor()
    
    if args.test:
        # Test with sample files
        test_files = [
            "README.md",
            "requirements.txt",
            "agents.json",
        ]
        
        for file_name in test_files:
            file_path = Path(file_name)
            if file_path.exists():
                print(f"\n{'='*60}")
                print(f"Testing: {file_name}")
                print('='*60)
                
                text, metadata = extractor.extract(file_path)
                
                print(f"Metadata: {json.dumps(metadata, indent=2)}")
                print(f"Text preview: {text[:200]}...")
                
        print(f"\nStats: {json.dumps(extractor.get_stats(), indent=2)}")
        
    elif args.file:
        file_path = Path(args.file)
        text, metadata = extractor.extract(file_path)
        
        print(f"File: {file_path}")
        print(f"Metadata: {json.dumps(metadata, indent=2)}")
        print(f"\nExtracted text:\n{text}")
        
    else:
        print("Usage:")
        print("  python universal_file_extractor.py <file>")
        print("  python universal_file_extractor.py --test")
        print("  python universal_file_extractor.py --install")